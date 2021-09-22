"""
Larry Barker
CS 521
8/17/18
Final Project
This module contains the Resume class with the following properties:
    An integer variable that holds the number of jobs found in the resume.
    An integer variable that holds the average length of jobs (in months).
    A list of the employment dates found in the resume.
    A method to get the docx file for the Resume.
    A method that returns the full text of the document.
    A method to get the word count in the resume.
    A method to get the name found in the resume.
    A method to count the number of jobs in the resume.
    A method that calculates and returns the average length of jobs (in months).
"""
import docx
import nltk, re

class Resume:
    numberOfJobs = 0
    lengthOfJobs = 0
    clean_dates = []
    
    def __init__(self, filename):
        self.filename = filename
        self.doc = docx.Document(filename)
        self.numberOfJobs = self.__countJobs(self.getText())
        self.lengthOfJobs = self.avgJobTime()

    def getDoc(self):       
        """
        Description: return a Document object for the resume
        :return: docx Document object from Resume filename
        """
        return docx.Document(self.filename)

    def getText(self):
        """
        Description: create a list that contains the text of each paragraph 
        in the document. Join each item in the list on a new line.
        :return: the joined list of text
        """
        fullText = []
        for para in self.doc.paragraphs:
            fullText.append(para.text)
        return '\n'.join(fullText)
    
    def getWordCount(self):
        """
        Description: call the getText() method and split the text to get a list
        of words found in the resume. Count the number of words as the length of
        the list.
        :return: wordCount = number of words in the resume
        """
        words = self.getText().split(" ")
        wordCount = len(words)
        return wordCount
    
    def getName(self):   
        """
        Description: Find the first paragraph in the document, which should be
        the person's name for the resume. This is used to merge resume data with
        employment data.
        :return: a concatenated string containing the name (Last name, First name)
        """
        firstParagraph = self.doc.paragraphs[0].text.split("\t")
        fullName = firstParagraph[0].split(" ")
        firstName = fullName[0]
        lastName = fullName[len(fullName) - 1]
        name = lastName + ", " + firstName
        return name
    
    def __countJobs(self, text):
        """
        Description: Count the number of jobs using Natural Language Processing.
        Each line the resume text is split, tokenized, and tagged with part of
        speech (POS). 

        Iterate over each line in the text, looking for the beginning of the 
        work experience section based on key words:
            experience, history, employment

        Save the index of the first line with any of the previous keywords.

        Perform a second iteration to find the end of the work experience
        section by looking for the next line that contains keyword(s):
            education, service, certificate
        and save the index.

        Slice the lines based on the indices found earlier and look for any

        :return: the length of the clean_dates list as the number of jobs
        """
        dates = []
        
        # split the text based on new lines
        lines = [el.strip() for el in text.split("\n") if len(el) > 0]
        
        # tokenize each line
        lines = [nltk.word_tokenize(el) for el in lines]
        
        # tag each token with part of speech
        lines = [nltk.pos_tag(el) for el in lines]
        

        # find the index of the sentence that begins work experience section
        for i, sentence in enumerate(lines):
            sen=" ".join([words[0].lower() for words in sentence])
            if len(sentence) <= 4 and (re.search('experience', sen) or re.search('history', sen) \
            or re.search('employment', sen)):
                startOfExperience = i
                break
            else:
                startOfExperience = 0
            
        # find the index of the sentence that ends work experience section
        for j, sentence in enumerate(lines):
            sen=" ".join([words[0].lower() for words in sentence])
            if (0 < len(sentence) <= 4) and (j > startOfExperience) \
            and (re.search('education', sen) or re.search('service', sen) \
                 or re.search('certificate', sen)):
                endOfExperience = j
                break
            else:
                endOfExperience = len(lines) - 1
        
        # slice the lines based on the start and end of experience
        for k, token in enumerate(lines[startOfExperience:endOfExperience]):
            # save the token if the POS is a cardinal digit, the length is > 2
            # and it is a number
            dates.append([tup[0] for tup in token if tup[1] == 'CD' \
                and len(tup[0]) > 2 and tup[0].isdigit()])
        
        # filter the list of dates to remove empty entries
        self.clean_dates = list(filter(None, dates))
        
        # look through dates for month/year or month-year formats 
        for n, date in enumerate(dates):
            for m, string in enumerate(date):
                # ignore strings that are alphabetic characters or less than
                # 3 characters
                if string.isalpha() or len(string) < 4:
                    date.remove(string)

                # find dates with format MONTH/YEAR and split
                if re.search('/', string):  
                    month_year = string.split('/')
                    date[m] = month_year[1] # save the year
        
        
        for m, date in enumerate(self.clean_dates):
            if len(date) == 0: # ignore empty dates
                continue
            # find dates with format month-year or year-year
            elif len(date[0]) > 4 and re.search('-', date[0]):
                self.clean_dates[m] = date[0].split('-')
            else:
                break
        
        return len(self.clean_dates)
    
    def avgJobTime(self):    
        """
        Description: Caclculate the average length of jobs (in months) by
        subracting the years. If no year is found, continue. If only 1 year is
        found, add 1 to the length of jobs. Otherwise, subtract the 1st date
        from the second. Divide the total length of jobs by the number of jobs
        and multiply by 12 to get the average months.
        :return: average = average length of jobs in months
        """
        for date in self.clean_dates:
            if len(date) == 0:
                continue
            elif len(date) == 1:
                self.lengthOfJobs += 1
            else:
                self.lengthOfJobs += int(date[1]) - int(date[0])
        if self.numberOfJobs == 0:
            return 0
        average = (self.lengthOfJobs / self.numberOfJobs) * 12
        return average